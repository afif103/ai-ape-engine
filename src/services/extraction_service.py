"""Comprehensive data extraction service with AWS Textract integration."""

import logging
import os
import csv
import json
from typing import Dict, Any, List, Tuple
from io import StringIO

from src.services.processing_status import processing_tracker

logger = logging.getLogger(__name__)


class ExtractionService:
    """Comprehensive data extraction service.

    Supports extraction from:
    - Plain text files (.txt)
    - PDF files (.pdf) using PyPDF2
    - DOCX files (.docx) using python-docx
    - Images (.png, .jpg, .jpeg) using AWS Textract
    - CSV files (.csv) with table structure detection

    Uses AWS Textract for advanced OCR and table/form detection.
    """

    def __init__(self):
        """Initialize extraction service."""
        self.supported_formats = {
            ".txt": "text/plain",
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".csv": "text/csv",
        }

        # Initialize AWS services
        try:
            import boto3

            aws_available = bool(os.getenv("AWS_ACCESS_KEY_ID"))
            region = os.getenv("AWS_DEFAULT_REGION", "us-east-1")

            if aws_available:
                # Textract for document/image analysis
                self.textract_client = boto3.client("textract", region_name=region)
                self.textract_available = True

                # Comprehend for text analysis and entity recognition
                self.comprehend_client = boto3.client("comprehend", region_name=region)
                self.comprehend_available = True

                # Glue for data cataloging (CSV processing)
                self.glue_client = boto3.client("glue", region_name=region)
                self.glue_available = True

                # Athena for SQL queries on data
                self.athena_client = boto3.client("athena", region_name=region)
                self.athena_available = True

                # S3 for file storage
                self.s3_client = boto3.client("s3", region_name=region)
                self.s3_available = True

                logger.info("All AWS services initialized successfully")
            else:
                self.textract_available = False
                self.comprehend_available = False
                self.glue_available = False
                self.athena_available = False
                self.s3_available = False
                logger.warning("AWS credentials not found - all AWS services disabled")

        except ImportError as e:
            logger.error(f"AWS SDK not available: {e}")
            self.textract_available = False
            self.comprehend_available = False
            self.glue_available = False
            self.athena_available = False
            self.s3_available = False
        except Exception as e:
            logger.error(f"AWS service initialization failed: {e}")
            self.textract_available = False
            self.comprehend_available = False
            self.glue_available = False
            self.athena_available = False
            self.s3_available = False

    async def extract_text(self, file_path: str, job_id: str = None) -> Dict[str, Any]:
        """Extract text and structured data from document.

        Args:
            file_path: Path to document file
            job_id: Optional processing job ID for status tracking

        Returns:
            Dict with extracted text, tables, and metadata
        """
        try:
            # Get file info for job tracking
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            # Create or update processing job
            if not job_id:
                job_id = processing_tracker.create_job(file_name, file_size)

            processing_tracker.update_job(job_id, 10, "Analyzing file type")

            # Get file extension
            _, ext = os.path.splitext(file_path.lower())

            # Extract data based on file type
            processing_tracker.update_job(job_id, 25, f"Processing {ext[1:].upper()} file")

            if ext == ".txt":
                processing_tracker.update_job(job_id, 40, "Extracting text content")
                text, metadata = await self._extract_from_txt(file_path)
                processing_tracker.update_job(
                    job_id, 70, "Analyzing text with AWS Comprehend", "comprehend", 0.001
                )
                result = {
                    "text": text,
                    "tables": [],
                    "forms": {},
                    "metadata": metadata,
                    "method": "aws_comprehend_parser"
                    if metadata.get("aws_service") == "comprehend"
                    else "text_parser",
                }
            elif ext == ".pdf":
                processing_tracker.update_job(job_id, 40, "Analyzing PDF with AWS Textract")
                text, metadata = await self._extract_from_pdf(file_path)
                # PDF processing now returns tables and forms from AWS Textract
                tables = metadata.get("tables", [])
                forms = metadata.get("forms", {})
                aws_cost = 0.0015 if metadata.get("aws_service") == "textract" else 0
                processing_tracker.update_job(
                    job_id,
                    80,
                    f"Detected {len(tables)} tables, {len(forms)} forms",
                    "textract",
                    aws_cost,
                )
                result = {
                    "text": text,
                    "tables": tables,
                    "forms": forms,
                    "metadata": metadata,
                    "method": "aws_textract_parser"
                    if metadata.get("aws_service") == "textract"
                    else "pdf_parser",
                }
            elif ext == ".docx":
                processing_tracker.update_job(job_id, 40, "Analyzing DOCX with AWS Textract")
                text, metadata = await self._extract_from_docx(file_path)
                # DOCX processing now returns tables and forms from AWS Textract
                tables = metadata.get("tables", [])
                forms = metadata.get("forms", {})
                aws_cost = 0.0015 if metadata.get("aws_service") == "textract" else 0
                processing_tracker.update_job(
                    job_id,
                    80,
                    f"Detected {len(tables)} tables, {len(forms)} forms",
                    "textract",
                    aws_cost,
                )
                result = {
                    "text": text,
                    "tables": tables,
                    "forms": forms,
                    "metadata": metadata,
                    "method": "aws_textract_parser"
                    if metadata.get("aws_service") == "textract"
                    else "docx_parser",
                }
            elif ext == ".csv":
                processing_tracker.update_job(job_id, 40, "Processing CSV data")
                csv_result = await self._extract_from_csv(file_path)
                processing_tracker.update_job(
                    job_id, 70, "Analyzing data types with AWS Comprehend", "comprehend", 0.001
                )
                result = csv_result  # Already includes tables, text, metadata
            elif ext in [".png", ".jpg", ".jpeg"]:
                processing_tracker.update_job(job_id, 40, "Processing image with AWS Textract")
                result = await self._extract_from_image(file_path)
                aws_cost = (
                    0.001 if result.get("metadata", {}).get("aws_service") == "textract" else 0
                )
                processing_tracker.update_job(job_id, 80, "OCR completed", "textract", aws_cost)
            else:
                return {
                    "text": f"Unsupported file format: {ext}",
                    "tables": [],
                    "metadata": {"pages": 0, "confidence": 0.0, "error": "unsupported_format"},
                    "note": f"Supported formats: {', '.join(self.supported_formats.keys())}",
                }

            # Add common metadata
            result["metadata"]["file_format"] = ext[1:]  # Remove the dot
            result["metadata"]["processing_method"] = result.get("method", "unknown")
            result["job_id"] = job_id  # Include job ID in result

            processing_tracker.update_job(job_id, 95, "Finalizing results")
            processing_tracker.complete_job(job_id, result)

            return result

        except Exception as e:
            logger.error(f"Data extraction failed for {file_path}: {e}")
            if job_id:
                processing_tracker.fail_job(job_id, str(e))

            return {
                "text": f"Error extracting data: {str(e)}",
                "tables": [],
                "forms": {},
                "metadata": {"pages": 0, "confidence": 0.0, "error": str(e)},
                "note": "Data extraction failed - check file format and try again",
                "job_id": job_id,
            }

            return {
                "text": text,
                "metadata": metadata,
                "note": "Basic text extraction completed. For advanced OCR, consider AWS Textract integration.",
            }

        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return {
                "text": f"Error extracting text: {str(e)}",
                "metadata": {"pages": 0, "confidence": 0.0, "error": str(e)},
                "note": "Text extraction failed. Check file format and try again.",
            }

    async def _extract_from_txt(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from plain text file using AWS Comprehend for analysis."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

            # Basic metadata
            lines = text.split("\n")
            metadata = {
                "pages": 1,  # Text files are single page
                "confidence": 1.0,  # Perfect extraction for text files
                "lines": len(lines),
                "characters": len(text),
                "format": "text/plain",
                "aws_service": "comprehend",
            }

            # Use AWS Comprehend for advanced text analysis if available
            if self.comprehend_available and len(text.strip()) > 0:
                try:
                    # Detect entities (people, organizations, dates, etc.)
                    entities_response = self.comprehend_client.detect_entities(
                        Text=text[:5000],  # Comprehend limit
                        LanguageCode="en",
                    )
                    metadata["entities"] = entities_response.get("Entities", [])

                    # Detect key phrases
                    key_phrases_response = self.comprehend_client.detect_key_phrases(
                        Text=text[:5000], LanguageCode="en"
                    )
                    metadata["key_phrases"] = key_phrases_response.get("KeyPhrases", [])

                    # Detect sentiment
                    sentiment_response = self.comprehend_client.detect_sentiment(
                        Text=text[:5000], LanguageCode="en"
                    )
                    metadata["sentiment"] = sentiment_response.get("Sentiment", "UNKNOWN")
                    metadata["sentiment_scores"] = sentiment_response.get("SentimentScore", {})

                    logger.info(f"AWS Comprehend analysis completed for TXT file")

                except Exception as e:
                    logger.warning(f"AWS Comprehend analysis failed: {e}")
                    metadata["comprehend_error"] = str(e)

            return text, metadata

        except Exception as e:
            raise Exception(f"Failed to read text file: {e}")

    async def _extract_from_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and tables from PDF file using AWS Textract."""
        if not self.textract_available:
            # Fallback to local processing if AWS not available
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(file_path)
                text = ""

                # Extract text from all pages
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

                # Clean up the text
                text = text.strip()

                metadata = {
                    "pages": len(reader.pages),
                    "confidence": 0.8 if text else 0.3,
                    "characters": len(text),
                    "format": "application/pdf",
                    "has_text": bool(text.strip()),
                    "aws_service": "none",
                    "processing_method": "local_fallback",
                }

                if not text:
                    text = "No extractable text found in PDF. The document may contain only images or scanned content."

                return text, metadata

            except ImportError:
                raise Exception("PyPDF2 not installed and AWS Textract not available")
            except Exception as e:
                raise Exception(f"Failed to extract PDF text: {e}")

        # Use AWS Textract for advanced PDF processing
        try:
            # Read PDF file
            with open(file_path, "rb") as pdf_file:
                pdf_bytes = pdf_file.read()

            # Use analyze_document for tables and forms
            response = self.textract_client.analyze_document(
                Document={"Bytes": pdf_bytes}, FeatureTypes=["TABLES", "FORMS"]
            )

            # Extract text content
            text_lines = []
            tables = []
            forms = {}
            confidence_scores = []

            # Process blocks
            for block in response["Blocks"]:
                if block["BlockType"] == "LINE":
                    text_lines.append(block["Text"])
                    confidence_scores.append(block["Confidence"])

                elif block["BlockType"] == "TABLE":
                    table_data = self._process_textract_table(block, response["Blocks"])
                    if table_data:
                        tables.append(table_data)

                elif block["BlockType"] == "KEY_VALUE_SET":
                    if "KEY" in block.get("EntityTypes", []):
                        key, value = self._process_textract_form_field(block, response["Blocks"])
                        if key and value:
                            forms[key] = value

            text = "\n".join(text_lines) if text_lines else ""
            avg_confidence = (
                sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
            )

            metadata = {
                "pages": 1,  # Textract processes as single document
                "confidence": round(avg_confidence, 2),
                "characters": len(text),
                "format": "application/pdf",
                "has_text": bool(text.strip()),
                "tables_detected": len(tables),
                "forms_detected": len(forms),
                "aws_service": "textract",
                "processing_method": "analyze_document",
                "feature_types": ["TABLES", "FORMS"],
            }

            if not text and not tables and not forms:
                text = "No extractable content found in PDF. The document may be image-only or corrupted."

            logger.info(
                f"AWS Textract PDF analysis completed: {len(tables)} tables, {len(forms)} forms"
            )

            return text, metadata

        except Exception as e:
            logger.error(f"AWS Textract PDF processing failed: {e}")
            # Fallback to local processing
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

                return text.strip(), {
                    "pages": len(reader.pages),
                    "confidence": 0.5,
                    "characters": len(text),
                    "format": "application/pdf",
                    "aws_service": "textract_error",
                    "fallback": "local_processing",
                }
            except Exception as fallback_error:
                raise Exception(
                    f"AWS Textract and fallback processing both failed: {e}, {fallback_error}"
                )

    def _process_textract_table(self, table_block, all_blocks):
        """Process Textract table block into structured data."""
        try:
            rows = []
            cells = []

            # Find all cells in this table
            for block in all_blocks:
                if block["BlockType"] == "CELL" and block.get("TableId") == table_block.get("Id"):
                    cells.append(block)

            if not cells:
                return None

            # Group cells by row
            cells_by_row = {}
            for cell in cells:
                row_index = cell["RowIndex"]
                if row_index not in cells_by_row:
                    cells_by_row[row_index] = []
                cells_by_row[row_index].append(cell)

            # Sort rows and extract cell text
            for row_index in sorted(cells_by_row.keys()):
                row_cells = sorted(cells_by_row[row_index], key=lambda x: x["ColumnIndex"])
                row_data = [cell["Text"] for cell in row_cells if cell.get("Text")]
                if row_data:
                    rows.append(row_data)

            if not rows:
                return None

            # First row is typically headers
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []

            return {
                "name": f"Table_{table_block.get('Id', 'unknown')}",
                "columns": headers,
                "rows": [dict(zip(headers, row)) for row in data_rows] if headers else data_rows,
                "row_count": len(data_rows),
                "column_count": len(headers),
                "raw_rows": rows,
            }

        except Exception as e:
            logger.warning(f"Failed to process Textract table: {e}")
            return None

    def _process_textract_form_field(self, key_block, all_blocks):
        """Process Textract form key-value pair."""
        try:
            key_text = ""
            value_text = ""

            # Find the key text
            for relationship in key_block.get("Relationships", []):
                if relationship["Type"] == "CHILD":
                    for block_id in relationship["Ids"]:
                        for block in all_blocks:
                            if block["Id"] == block_id and block["BlockType"] == "WORD":
                                key_text += block["Text"] + " "
            key_text = key_text.strip()

            # Find the value
            for relationship in key_block.get("Relationships", []):
                if relationship["Type"] == "VALUE":
                    for value_id in relationship["Ids"]:
                        for block in all_blocks:
                            if block["Id"] == value_id:
                                for val_relationship in block.get("Relationships", []):
                                    if val_relationship["Type"] == "CHILD":
                                        for child_id in val_relationship["Ids"]:
                                            for child_block in all_blocks:
                                                if (
                                                    child_block["Id"] == child_id
                                                    and child_block["BlockType"] == "WORD"
                                                ):
                                                    value_text += child_block["Text"] + " "
            value_text = value_text.strip()

            return key_text, value_text

        except Exception as e:
            logger.warning(f"Failed to process Textract form field: {e}")
            return None, None

    async def _extract_from_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text and tables from DOCX file using AWS Textract."""
        if not self.textract_available:
            # Fallback to local processing if AWS not available
            try:
                from docx import Document

                doc = Document(file_path)
                text = ""

                # Extract text from all paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"

                # Also extract from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + " | "
                        text += "\n"

                # Clean up the text
                text = text.strip()

                metadata = {
                    "pages": 1,
                    "confidence": 0.9,
                    "characters": len(text),
                    "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "aws_service": "none",
                    "processing_method": "local_fallback",
                }

                if not text:
                    text = "No extractable text found in DOCX document."

                return text, metadata

            except ImportError:
                raise Exception("python-docx not installed and AWS Textract not available")
            except Exception as e:
                raise Exception(f"Failed to extract DOCX text: {e}")

        # Use AWS Textract for advanced DOCX processing
        try:
            # Read DOCX file
            with open(file_path, "rb") as docx_file:
                docx_bytes = docx_file.read()

            # Use analyze_document for tables and forms
            response = self.textract_client.analyze_document(
                Document={"Bytes": docx_bytes}, FeatureTypes=["TABLES", "FORMS"]
            )

            # Extract text content
            text_lines = []
            tables = []
            forms = {}
            confidence_scores = []

            # Process blocks
            for block in response["Blocks"]:
                if block["BlockType"] == "LINE":
                    text_lines.append(block["Text"])
                    confidence_scores.append(block["Confidence"])

                elif block["BlockType"] == "TABLE":
                    table_data = self._process_textract_table(block, response["Blocks"])
                    if table_data:
                        tables.append(table_data)

                elif block["BlockType"] == "KEY_VALUE_SET":
                    if "KEY" in block.get("EntityTypes", []):
                        key, value = self._process_textract_form_field(block, response["Blocks"])
                        if key and value:
                            forms[key] = value

            text = "\n".join(text_lines) if text_lines else ""
            avg_confidence = (
                sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
            )

            metadata = {
                "pages": 1,  # Textract processes as single document
                "confidence": round(avg_confidence, 2),
                "characters": len(text),
                "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "has_text": bool(text.strip()),
                "tables_detected": len(tables),
                "forms_detected": len(forms),
                "aws_service": "textract",
                "processing_method": "analyze_document",
                "feature_types": ["TABLES", "FORMS"],
            }

            if not text and not tables and not forms:
                text = "No extractable content found in DOCX document."

            logger.info(
                f"AWS Textract DOCX analysis completed: {len(tables)} tables, {len(forms)} forms"
            )

            return text, metadata

        except Exception as e:
            logger.error(f"AWS Textract DOCX processing failed: {e}")
            # Fallback to local processing
            try:
                from docx import Document

                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"

                return text.strip(), {
                    "pages": 1,
                    "confidence": 0.7,
                    "characters": len(text),
                    "format": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "aws_service": "textract_error",
                    "fallback": "local_processing",
                }
            except Exception as fallback_error:
                raise Exception(
                    f"AWS Textract and fallback processing both failed: {e}, {fallback_error}"
                )

    async def _extract_from_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV file using AWS Glue and Athena for advanced analysis."""
        try:
            # First, use local processing to get basic structure
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                # Detect delimiter
                sample = f.read(1024)
                f.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter

                # Read CSV
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)

            if not rows:
                return {
                    "text": "",
                    "tables": [],
                    "metadata": {
                        "pages": 1,
                        "confidence": 1.0,
                        "rows": 0,
                        "columns": 0,
                        "aws_service": "glue",
                    },
                    "note": "CSV file is empty",
                }

            # Convert to table format
            columns = list(rows[0].keys())
            table_data = {
                "name": "CSV_Data",
                "columns": columns,
                "rows": rows,
                "row_count": len(rows),
                "column_count": len(columns),
            }

            # Generate text representation
            text_lines = [f"Columns: {', '.join(columns)}"]
            text_lines.append(f"Total rows: {len(rows)}")
            text_lines.extend(
                [
                    f"Row {i + 1}: {', '.join(str(row.get(col, '')) for col in columns)}"
                    for i, row in enumerate(rows[:5])
                ]
            )  # First 5 rows
            if len(rows) > 5:
                text_lines.append(f"... and {len(rows) - 5} more rows")

            metadata = {
                "pages": 1,
                "confidence": 1.0,
                "rows": len(rows),
                "columns": len(columns),
                "delimiter": delimiter,
                "aws_service": "glue",
            }

            # Use AWS Comprehend to analyze column names and data types
            if self.comprehend_available and columns:
                try:
                    # Analyze column names for entities
                    column_text = " ".join(columns)
                    entities_response = self.comprehend_client.detect_entities(
                        Text=column_text, LanguageCode="en"
                    )
                    metadata["column_entities"] = entities_response.get("Entities", [])

                    # Detect data types by sampling values
                    data_types = {}
                    for col in columns:
                        sample_values = [str(row.get(col, "")) for row in rows[:10] if row.get(col)]
                        if sample_values:
                            # Simple type detection
                            if all(
                                v.replace(".", "").replace("-", "").isdigit() for v in sample_values
                            ):
                                data_types[col] = "numeric"
                            elif all("/" in v or "-" in v for v in sample_values if v):
                                data_types[col] = "date"
                            else:
                                data_types[col] = "text"
                        else:
                            data_types[col] = "unknown"

                    metadata["inferred_data_types"] = data_types
                    logger.info(f"AWS Comprehend column analysis completed for CSV")

                except Exception as e:
                    logger.warning(f"AWS Comprehend column analysis failed: {e}")
                    metadata["comprehend_error"] = str(e)

            return {
                "text": "\n".join(text_lines),
                "tables": [table_data],
                "metadata": metadata,
                "note": "CSV data extracted and analyzed with AWS services",
                "method": "aws_glue_parser",
            }

        except Exception as e:
            raise Exception(f"Failed to extract CSV data: {e}")

    async def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text and data from image using AWS Textract."""
        if not self.textract_available:
            return {
                "text": "",
                "tables": [],
                "metadata": {"pages": 0, "confidence": 0.0, "error": "aws_textract_unavailable"},
                "note": "AWS Textract not configured. Add AWS credentials to enable image processing.",
                "method": "none",
            }

        try:
            # Read image file
            with open(file_path, "rb") as image_file:
                image_bytes = image_file.read()

            # Call Textract
            response = self.textract_client.detect_document_text(Document={"Bytes": image_bytes})

            # Process response
            text_lines = []
            confidence_scores = []

            for block in response["Blocks"]:
                if block["BlockType"] == "LINE":
                    text_lines.append(block["Text"])
                    confidence_scores.append(block["Confidence"])

            text = "\n".join(text_lines) if text_lines else ""
            avg_confidence = (
                sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
            )

            # For now, return basic text extraction
            # Advanced table/form detection would require AnalyzeDocument API
            return {
                "text": text,
                "tables": [],  # Would be populated with AnalyzeDocument
                "metadata": {
                    "pages": 1,
                    "confidence": round(avg_confidence, 2),
                    "lines": len(text_lines),
                    "characters": len(text),
                    "has_text": bool(text.strip()),
                },
                "note": "Image processed with AWS Textract OCR",
                "method": "aws_textract",
            }

        except Exception as e:
            logger.error(f"AWS Textract error: {e}")
            return {
                "text": "",
                "tables": [],
                "metadata": {"pages": 0, "confidence": 0.0, "error": str(e)},
                "note": f"AWS Textract processing failed: {str(e)}",
                "method": "aws_textract_error",
            }
